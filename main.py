from selenium.webdriver.common.action_chains import ActionChains
from docx import Document
from selenium import webdriver
from datetime import date


class Paragraph:
    def __init__(self, key_words, content, company):
        self.key_words = key_words
        self.content = content
        self.company = company
        self.compatibility = 0


def get_num_jobs():
    jobs = int(driver.find_element_by_xpath('//*[@id="postingsTablePlaceholder"]/div[1]/span[1]').text)
    return jobs


def login():
    student = driver.find_element_by_xpath('/html/body/div[3]/div/div/div[5]/div/div/a[1]')
    student.click()

    email = driver.find_element_by_xpath('//*[@id="userNameInput"]')
    email.send_keys('EMAILHERE')

    enteremail = driver.find_element_by_xpath('//*[@id="nextButton"]')
    enteremail.click()


def find_job(job_num):
    actions = ActionChains(driver)

    button = driver.find_element_by_xpath('//*[@id="postingsTable"]/tbody/tr[' + str(job_num) + ']/td[1]/div/a')
    actions.move_to_element(button).perform()
    button.click()
    no_new_tab = driver.find_element_by_xpath(
        '//*[@id="postingsTable"]/tbody/tr[' + str(job_num) + ']/td[1]/div/ul/li[1]/a')
    no_new_tab.click()


def find_keywords(responsibilities, key_words, skills):

    counter = [0] * len(key_words)

    resp_words = responsibilities.text.split()
    skills_words = skills.text.split()
    for i in resp_words:
        for k in range(len(key_words)):
            if i.replace(",", "").replace(".", "").lower() == key_words[k]:
                counter[k] += 1
    for i in skills_words:
        for k in range(len(key_words)):
            if i.replace(",", "").replace(".", "").lower() == key_words[k]:
                counter[k] += 1
    for i in range(len(key_words)):
        print(key_words[i] + ": " + str(counter[i]))
    return counter


def write_paragraph(file, company, key_words, counter):
    read_file = open(file, "r")
    num_intro = int(read_file.readline())
    paragraphs = []
    for i in range(num_intro):
        words = read_file.readline().split()
        paragraph = read_file.readline()
        paragraphs.append(Paragraph(words, paragraph, company))
    best_paragraph = "[NEED HELP]"
    highest_compat_num = 0
    for i in range(num_intro):
        for k in range(len(key_words)):
            for j in paragraphs[i].key_words:
                if j == key_words[k]:
                    paragraphs[i].compatibility += counter[k]
                    if paragraphs[i].compatibility > highest_compat_num:
                        highest_compat_num = paragraphs[i].compatibility
                        best_paragraph = paragraphs[i].content
    best_paragraph = best_paragraph.replace("@@@@@@", company)
    best_paragraph = best_paragraph.replace("..", ".")
    read_file.close()
    return " ".join(best_paragraph.split())


def get_data():
    actions = ActionChains(driver)

    skills = responsibilities = division = driver.find_element_by_xpath(
        '//*[@id="postingDiv"]/div[3]/div[2]/table/tbody/tr[2]/td[2]')
    jobid = driver.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div[1]/div[1]/h1')
    temp = ([int(i) for i in jobid.text.split() if i.isdigit()])
    id = "#"
    for i in temp:
        id += str(i)
    company = driver.find_element_by_xpath('//*[@id="postingDiv"]/div[3]/div[2]/table/tbody/tr[1]/td[2]')
    jobtitle = street = city = province = postalcode = \
        driver.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div[1]/div[1]/h1')

    rows = driver.find_elements_by_xpath('//*[@id="postingDiv"]/div[1]/div[2]/table/tbody/tr/td[1]')
    num = 1
    for i in rows:
        row = driver.find_element_by_xpath('//*[@id="postingDiv"]/div[1]/div[2]/table/tbody/tr[' + str(num) + ']/td[2]')
        actions.move_to_element(row).perform()
        print(i.text)
        if i.text.lower() == "job title:":
            jobtitle = row
        elif i.text.lower() == "job - address line one:":
            street = row
        elif i.text.lower() == "job - city:":
            city = row
        elif i.text.lower() == "job - province / state:":
            province = row
        elif i.text.lower() == "job - postal code / zip code (x#x #x#):":
            postalcode = row
        elif i.text.lower() == "job responsibilities:":
            responsibilities = row
        elif i.text.lower() == "required skills:":
            skills = row
        num += 1

    write_cover_letter(id, jobtitle, division, company, street, city, province, postalcode, responsibilities, skills)
    go_back = driver.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div[1]/div[2]/ul/li[2]/a')
    actions.move_to_element(go_back).perform()
    go_back.click()


def write_cover_letter(id, job_title, division, company, street, city, province, postalcode, responsiblities, skills):
    document = Document("Template.docx")

    font = document.styles['Normal'].font
    font.name = "TEXT STYLE"

    current_day = date.today()
    current_day = current_day.strftime("%b %d, %Y")

    #files to look in
    intro = "intro.txt"
    p1 = "p1.txt"
    p2 = "p2.txt"
    key_words = ["KEYWORDS"]
    counter = find_keywords(responsiblities, key_words, skills)

    document.paragraphs[0].add_run(current_day)
    document.add_paragraph()
    document.add_paragraph(division.text)
    document.add_paragraph(company.text)
    document.add_paragraph(street.text)
    document.add_paragraph(city.text + ", " + province.text + ", " + postalcode.text)
    document.add_paragraph()
    document.add_paragraph("Dear Hiring Manager,")
    document.add_paragraph("These are just some of the skills I possess that would make me a great fit for "
                           + company.text + ". I would love to answer any questions or talk more about my experiences"
                                            " in an interview. Thank you for your time.")
    document.add_paragraph()
    document.add_paragraph("Sincerely,")
    document.add_paragraph("NAME")

    document.save("C:/Users/bobco/Documents/Resume Stuff/Cover Letters/" + id + ".docx")


driver = webdriver.Chrome(executable_path='C:/Users/bobco/Documents/chromedriver_win32/chromedriver')
driver.get("https://waterlooworks.uwaterloo.ca")

login()

go = input("Write cover letter?")

num_jobs = 20  # get_num_jobs()
current_job = 1

while go != 'n' and current_job <= num_jobs:
    # find_job(current_job)
    get_data()
    # current_job += 1
    go = input("Next job?")

quit()
